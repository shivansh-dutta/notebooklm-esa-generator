"""
scripts/export_docx.py — Phase 1 ESA Report Generator

Fills the real Envicon_Phase_I_ESA_Report_TEMPLATE.docx in place: substitutes
`{{placeholder}}` tokens with project metadata, injects each drafted
subsection's prose immediately after its matching template heading, and
(Phase B) populates the template's real tables from pipeline data.

This replaces the old Pandoc-based DOCX export in scripts/export.py, which
pointed at a `TemplateVault/pandoc/reference.docx` that never existed and so
silently fell back to unstyled default output. The template itself is now
the single source of truth for headings, numbering, styles, cover, the live
Table of Contents, and the Signature Page — nothing about its fixed
structure is regenerated from scratch.

Public interface:
    run_export_docx(project_path: Path) -> Path
"""

from __future__ import annotations

import json
import logging
import re
from pathlib import Path

from docx import Document

from scripts.docx_helpers import (
    append_row_like,
    build_heading_index,
    enable_update_fields_on_open,
    find_table_after_heading,
    find_table_by_header,
    find_table_by_header_and_first_row,
    insert_blocks_after,
    markdown_lite_to_blocks,
    normalize_heading,
    remove_paragraphs_after_heading_matching,
    remove_paragraphs_matching,
    replace_placeholders,
    set_cell_text,
)
from scripts.report_constants import (
    DATABASE_TO_LIST,
    EDR_TABLE_HEADER,
    FEDERAL_LIST_FIRST_ROW,
    PLACEHOLDER_FIELD_MAP,
    STATE_LIST_FIRST_ROW,
    pe_marker,
)

logger = logging.getLogger(__name__)

REPO_ROOT = Path(__file__).resolve().parent.parent
TEMPLATE_PATH = REPO_ROOT / "Envicon_Phase_I_ESA_Report_TEMPLATE.docx"

_FRONTMATTER_RE = re.compile(r"^---\s*\n(.*?)\n---\s*\n?", re.DOTALL)
_FIELD_RE = re.compile(r"^(\w[\w_]*):\s*(.+)$", re.MULTILINE)
_SUBSECTION_HEADING_RE = re.compile(r"^(#{1,4})\s+(.+?)\s*$", re.MULTILINE)


# ---------------------------------------------------------------------------
# Project context loading
# ---------------------------------------------------------------------------

def _parse_frontmatter_block(text: str) -> dict[str, str]:
    """Parse a leading `---\\n key: value \\n---` YAML-ish frontmatter block
    into a flat {key: value} dict (quotes stripped). Returns {} if no
    frontmatter block is found."""
    m = _FRONTMATTER_RE.match(text)
    if not m:
        return {}
    return {k: v.strip().strip('"') for k, v in _FIELD_RE.findall(m.group(1))}


def load_dashboard_meta(project_path: Path) -> dict[str, str]:
    """Parse 00_Project_Dashboard.md frontmatter; return {} if absent/unreadable."""
    dash = project_path / "00_Project_Dashboard.md"
    if not dash.exists():
        return {}
    try:
        text = dash.read_text(encoding="utf-8")
    except OSError:
        return {}
    return _parse_frontmatter_block(text)


def load_edr_hit_records(project_path: Path) -> list[dict[str, str]]:
    """
    Parse the frontmatter of every EDR hit note in both
    <project>/EDR_Database_Hits/ (within the 0.1-mi auto-draft radius) and
    <project>/Manual_Review/ (beyond that radius, or unparseable distance).

    Both folders are read here — the Researcher's 0.1-mi routing radius
    controls auto-drafting eligibility, not the much larger ASTM Table 1
    list search radii (0.25-1.0 mi) that the template's EDR radius tables
    report against, so a record routed to Manual_Review/ can still fall
    within a list's statutory search radius and must still be counted.
    """
    records: list[dict[str, str]] = []
    skip = {"_index.md", "no_hits.md"}
    for folder_name in ("EDR_Database_Hits", "Manual_Review"):
        folder = project_path / folder_name
        if not folder.exists():
            continue
        for md_file in sorted(folder.glob("*.md")):
            if md_file.name in skip:
                continue
            try:
                text = md_file.read_text(encoding="utf-8")
            except OSError:
                continue
            fields = _parse_frontmatter_block(text)
            if fields:
                fields["_source_folder"] = folder_name
                records.append(fields)
    return records


def strip_frontmatter(text: str) -> str:
    """Remove a leading YAML frontmatter block (--- ... ---), if present."""
    m = _FRONTMATTER_RE.match(text)
    return text[m.end():] if m else text


def parse_writer_sections(report_sections_dir: Path) -> dict[str, str]:
    """
    Read every .md file in *report_sections_dir* and split each into
    {normalized_heading: prose} blocks on markdown headings (# through ####).

    Content appearing before the first heading in a file is keyed under the
    file's own title — the first `# ` line if present, else a normalized
    version of the filename — so files with no internal subsection tags
    (e.g. the current 8-file model, prior to the Phase C 12-section rewrite)
    still contribute their content as a single block rather than being
    silently dropped. inject_sections() below simply skips any key that
    doesn't match a template heading, so this degrades gracefully instead of
    crashing when Writer output doesn't yet match the template's exact
    subsection structure.
    """
    if not report_sections_dir.exists():
        return {}

    sections: dict[str, str] = {}
    for md_file in sorted(report_sections_dir.glob("*.md")):
        try:
            text = strip_frontmatter(md_file.read_text(encoding="utf-8"))
        except OSError:
            continue

        matches = list(_SUBSECTION_HEADING_RE.finditer(text))
        if not matches:
            fallback_key = normalize_heading(md_file.stem.split("_", 1)[-1].replace("_", " "))
            if text.strip():
                sections[fallback_key] = text.strip()
            continue

        # Content before the first heading (if any) goes under a fallback key.
        preamble = text[: matches[0].start()].strip()
        if preamble:
            fallback_key = normalize_heading(md_file.stem.split("_", 1)[-1].replace("_", " "))
            sections[fallback_key] = preamble

        for i, m in enumerate(matches):
            heading_text = m.group(2)
            start = m.end()
            end = matches[i + 1].start() if i + 1 < len(matches) else len(text)
            body = text[start:end].strip()
            if body:
                sections[normalize_heading(heading_text)] = body

    return sections


# ---------------------------------------------------------------------------
# Injection + placeholder mapping
# ---------------------------------------------------------------------------

def inject_sections(doc, headings: dict, parsed: dict[str, str]) -> set[str]:
    """
    Insert each parsed subsection's prose immediately after its matching
    template heading. Returns the set of parsed keys that had no matching
    template heading (logged by the caller, not fatal — a section the
    Writer hasn't been updated to emit yet just leaves that template
    heading with no body).
    """
    unmatched: set[str] = set()
    for key, prose in parsed.items():
        heading_para = headings.get(key)
        if heading_para is None:
            unmatched.add(key)
            continue
        blocks = markdown_lite_to_blocks(prose)
        if blocks:
            insert_blocks_after(heading_para, blocks)
    return unmatched


def aggregate_edr_counts(records: list[dict[str, str]]) -> dict[str, int]:
    """
    Group hit-note records by their mapped ASTM Table 1 list name
    (DATABASE_TO_LIST) and count them. Records whose database_source has no
    mapping are skipped (logged, not counted, not guessed) — see
    report_constants.DATABASE_TO_LIST for why.
    """
    counts: dict[str, int] = {}
    unmapped: set[str] = set()
    for record in records:
        db_source = record.get("database_source", "")
        list_name = DATABASE_TO_LIST.get(db_source)
        if list_name is None:
            if db_source:
                unmapped.add(db_source)
            continue
        counts[list_name] = counts.get(list_name, 0) + 1

    if unmapped:
        logger.info(
            "export_docx: %d database_source value(s) have no ASTM list "
            "mapping and were not counted in the EDR radius tables "
            "(hit notes remain visible in EDR_Database_Hits/Manual_Review): %s",
            len(unmapped), sorted(unmapped),
        )
    return counts


def populate_edr_tables(doc, project_path: Path) -> None:
    """
    Fill the "Listings in radius" column of the template's Federal and
    State/Tribal/Local EDR radius tables (section 5.3) with real counts
    aggregated from this project's EDR hit notes.

    Only rows whose List name has a DATABASE_TO_LIST mapping are touched;
    all other cells (including "On Subject Property" and "REC relative to
    Subject Property", which require professional judgment the pipeline
    cannot supply) are left as their original `{{...}}` placeholder so the
    later replace_placeholders() pass marks them PE_MARKER — this function
    must run BEFORE that pass.
    """
    counts = aggregate_edr_counts(load_edr_hit_records(project_path))

    federal_table = find_table_by_header_and_first_row(
        doc, EDR_TABLE_HEADER, FEDERAL_LIST_FIRST_ROW
    )
    state_table = find_table_by_header_and_first_row(
        doc, EDR_TABLE_HEADER, STATE_LIST_FIRST_ROW
    )

    for table in (federal_table, state_table):
        if table is None:
            continue
        for row in table.rows[1:]:
            list_name = row.cells[0].text.strip()
            if list_name not in counts:
                continue  # no mapped hits found for this list — leave {{n}} for PE_MARKER
            set_cell_text(row.cells[3], str(counts[list_name]))


# ---------------------------------------------------------------------------
# Template scaffolding cleanup — literal paragraphs the template ships with
# that should never survive into an issued report (confirmed present in the
# real Envicon_Phase_I_ESA_Report_TEMPLATE.docx, not something the Writer or
# NotebookLM introduces).
# ---------------------------------------------------------------------------

_WRITER_NOTE_MARKER = "Writer note (delete before issue):"


def strip_writer_notes(doc) -> int:
    """
    Remove every literal "» Writer note (delete before issue): ..." paragraph
    baked into the template (13 confirmed occurrences — Revision History,
    Acronyms guidance, Executive Summary REC numbering, section 2/3/4/5/8
    guidance, the title-records fallback line, and the signature-page
    reminder). These are instructions to whoever fills the template in by
    hand, not report content, and nothing previously stripped them — they
    rode straight through into the exported DOCX. Returns the count removed.
    """
    return remove_paragraphs_matching(doc, lambda text: _WRITER_NOTE_MARKER in text)


# The template's own hardcoded "12.0 References" list (paragraphs that ship
# in Envicon_Phase_I_ESA_Report_TEMPLATE.docx itself, confirmed via direct
# inspection) — every one of these is a `{{placeholder}}` bullet that
# degrades to PE_MARKER text once replace_placeholders() runs, sitting
# directly beneath whatever real References list the Writer/NotebookLM
# drafted and inject_sections() already inserted right after the same
# heading. Matched by exact known substrings (not a heading-range scan) so
# this only ever removes these specific known lines, never real content.
_TEMPLATE_REFERENCES_BOILERPLATE = (
    "The following sources were reviewed as part of this assessment:",
    "ASTM International. Standard Practice for Environmental Site Assessments",
    "{{Database Provider}} Radius Map Report",
    "{{Database Provider}} Certified Sanborn Map Report",
    "{{Database Provider}} City Directory Package",
    "{{Database Provider}} Aerial Photographs",
    "{{Prior report title, author, and date}}",
    "{{Agency records reviewed",
)


def strip_duplicate_references_boilerplate(doc, parsed_sections: dict[str, str], headings: dict) -> int:
    """
    Remove the template's own placeholder References list ONLY when a real
    References section was actually drafted and injected under the "12.0
    References" heading (parsed_sections has that key) — if nothing was
    drafted, leave the template's placeholder list in place rather than
    deleting the section's only content. Must run AFTER inject_sections()
    (so parsed_sections reflects what was actually injected) and BEFORE
    replace_placeholders() (its {{...}} tokens must still be literal text to
    match against).

    Scoped to paragraphs after the "12.0 References" heading (stopping at
    the next heading) rather than a document-wide substring match — the
    `{{Database Provider}}` token this boilerplate uses is NOT unique to
    this section (Section 5.3's own intro paragraph reuses it), so a
    document-wide match would delete real content elsewhere.
    """
    heading_para = headings.get("12.0 references")
    if "12.0 references" not in parsed_sections or heading_para is None:
        return 0
    return remove_paragraphs_after_heading_matching(
        doc, heading_para, lambda text: any(marker in text for marker in _TEMPLATE_REFERENCES_BOILERPLATE)
    )


# ---------------------------------------------------------------------------
# Historical tables (aerial photos / Sanborn maps / city directories)
# ---------------------------------------------------------------------------

# (heading_key, historical_tables.json key, column order) — column order
# must match notebooklm_pipeline.question_bank.historical_table_questions()'s
# JSON schema.
_HISTORICAL_TABLE_SPECS = (
    ("5.2.1 aerial photographs", "aerial", ("year", "subject_property", "adjacent_properties")),
    ("5.2.2 sanborn fire insurance maps", "sanborn", ("year", "subject_property", "adjacent_properties")),
    ("5.2.3 city and street directories", "city_directory", ("year", "address", "occupant")),
)


def populate_historical_tables(doc, project_path: Path, headings: dict) -> None:
    """
    Fill the Aerial Photographs (5.2.1), Sanborn Maps (5.2.2), and City/
    Street Directories (5.2.3) tables from
    <project>/Historical_Records/historical_tables.json (written by
    notebooklm_pipeline.assemble.write_historical_tables — see
    notebooklm_pipeline.question_bank.historical_table_questions() for how
    that data is gathered).

    The Aerial and Sanborn tables share an IDENTICAL header row AND first
    data row (both "Year | Subject Property | Adjacent properties" with
    {{year}}/{{observation}}/{{observation}} placeholders), so
    find_table_by_header_and_first_row (used for the EDR tables) can't tell
    them apart — these are located by heading proximity instead
    (find_table_after_heading), which the real template's structure supports
    (each table sits immediately after its own numbered subsection heading,
    confirmed by direct inspection).

    A table with no data for its key is left with its original {{...}}
    placeholder row untouched, so the later replace_placeholders() pass
    marks it PE_MARKER — same convention as populate_edr_tables. Must run
    BEFORE that pass.
    """
    data_path = Path(project_path) / "Historical_Records" / "historical_tables.json"
    if not data_path.exists():
        return
    try:
        tables_data = json.loads(data_path.read_text(encoding="utf-8"))
    except (OSError, ValueError):
        logger.warning("export_docx: could not read/parse %s — skipping historical tables", data_path)
        return

    for heading_key, data_key, columns in _HISTORICAL_TABLE_SPECS:
        rows = tables_data.get(data_key) or []
        if not rows:
            continue
        heading_para = headings.get(heading_key)
        if heading_para is None:
            continue
        table = find_table_after_heading(doc, heading_para)
        if table is None or len(table.rows) < 2:
            continue

        # Fill every pre-existing placeholder row (below the header) before
        # appending new ones — filling only row[1] and leaving any further
        # template placeholder rows (e.g. a 2-example-row template) untouched
        # left a stale "{{...}}"->PE_MARKER row wedged between real data rows.
        existing_placeholder_rows = table.rows[1:]
        for idx, row_data in enumerate(rows):
            values = [str(row_data.get(c, "")) for c in columns]
            if idx < len(existing_placeholder_rows):
                for cell, value in zip(existing_placeholder_rows[idx].cells, values):
                    set_cell_text(cell, value)
            else:
                append_row_like(table, 1, values)

        # Remove any placeholder rows left over beyond what real data filled.
        for row in existing_placeholder_rows[len(rows):]:
            row._element.getparent().remove(row._element)


# ---------------------------------------------------------------------------
# Acronyms — prune rows for acronyms the finished document never actually
# uses (report_constants.py's docstring flagged this as a never-implemented
# "Phase B" idea; the 631 Northland review flagged the static 33-row table
# as never scanned against the finished doc).
# ---------------------------------------------------------------------------

_ACRONYM_TABLE_HEADER = ("Acronym", "Definition")


def prune_unused_acronyms(doc) -> int:
    """
    Remove rows from the Acronyms table whose acronym does not appear
    (whole-word, case-sensitive — acronyms are always written in caps)
    anywhere else in the finished document body. Must run LAST, after every
    other injection/placeholder pass, since it scans the document's final
    text. Returns the count of rows removed.

    Excludes the acronym table's own cells from the "is it used elsewhere"
    scan by re-matching each table's header row (same signature check
    find_table_by_header uses) during a single fresh doc.tables pass, rather
    than comparing paragraph/element objects across two separate traversals
    — python-docx/lxml don't guarantee a stable id() for element proxy
    objects re-materialized on each .paragraphs/.tables access, so an
    identity-based exclusion set can silently collide with unrelated
    elements once the original proxies are garbage-collected.
    """
    table = find_table_by_header(doc, _ACRONYM_TABLE_HEADER)
    if table is None or len(table.rows) < 2:
        return 0

    normalized_target = tuple(normalize_heading(h) for h in _ACRONYM_TABLE_HEADER)

    body_text_parts = [p.text for p in doc.paragraphs]
    for t in doc.tables:
        header_cells = tuple(normalize_heading(c.text) for c in t.rows[0].cells) if t.rows else ()
        if header_cells == normalized_target:
            continue  # this IS the acronym table — exclude its own cells
        for row in t.rows:
            body_text_parts.extend(cell.text for cell in row.cells)
    for section in doc.sections:
        for hf in (section.header, section.footer, section.first_page_header, section.first_page_footer):
            if hf is None:
                continue
            body_text_parts.extend(p.text for p in hf.paragraphs)

    body_text = "\n".join(body_text_parts)

    rows_to_remove = []
    for row in table.rows[1:]:
        acronym = row.cells[0].text.strip()
        if acronym and not re.search(rf"\b{re.escape(acronym)}\b", body_text):
            rows_to_remove.append(row)

    for row in rows_to_remove:
        row._tr.getparent().remove(row._tr)
    return len(rows_to_remove)


def build_placeholder_map(dashboard_meta: dict[str, str]) -> dict[str, str]:
    """
    Build the {normalized_placeholder_name: value} map used by
    replace_placeholders(), from PLACEHOLDER_FIELD_MAP + dashboard
    frontmatter. A dashboard field that is missing, empty, or the literal
    "TBD" placeholder value is left out of the map entirely — the exporter
    then substitutes the PE-completion marker for it rather than writing
    "TBD" or leaving the raw `{{...}}` token behind.
    """
    mapping: dict[str, str] = {}
    for placeholder_key, field_name in PLACEHOLDER_FIELD_MAP.items():
        value = dashboard_meta.get(field_name, "").strip()
        if value and value.upper() != "TBD":
            mapping[placeholder_key] = value
    return mapping


# ---------------------------------------------------------------------------
# Main entry point
# ---------------------------------------------------------------------------

def run_export_docx(project_path: Path, template_path: Path = TEMPLATE_PATH) -> Path:
    """
    Fill the Envicon Phase I ESA report template with this project's drafted
    content and metadata, and save it to
    <project_path>/Export/<project_name>_Phase1_ESA_DRAFT.docx.
    """
    project_path = Path(project_path).resolve()
    project_name = project_path.name

    if not template_path.exists():
        raise FileNotFoundError(f"Report template not found at {template_path}")

    export_dir = project_path / "Export"
    export_dir.mkdir(exist_ok=True)

    dashboard_meta = load_dashboard_meta(project_path)
    parsed_sections = parse_writer_sections(project_path / "Report_Sections")

    doc = Document(str(template_path))
    headings = build_heading_index(doc)

    removed_notes = strip_writer_notes(doc)
    if removed_notes:
        logger.info("export_docx: stripped %d template 'Writer note' paragraph(s)", removed_notes)

    unmatched = inject_sections(doc, headings, parsed_sections)
    if unmatched:
        logger.info(
            "export_docx: %d drafted section(s) had no matching template "
            "heading and were not inserted: %s",
            len(unmatched), sorted(unmatched),
        )

    removed_refs = strip_duplicate_references_boilerplate(doc, parsed_sections, headings)
    if removed_refs:
        logger.info(
            "export_docx: removed %d line(s) of the template's duplicate "
            "placeholder References list (real References were drafted)",
            removed_refs,
        )

    # Both must run before replace_placeholders(): they fill specific cells
    # with real per-row data, leaving everything else's `{{...}}` untouched
    # for the generic placeholder pass to mark PE_MARKER.
    populate_historical_tables(doc, project_path, headings)
    populate_edr_tables(doc, project_path)

    placeholder_map = build_placeholder_map(dashboard_meta)
    unresolved = replace_placeholders(doc, placeholder_map, unresolved_marker=pe_marker())
    if unresolved:
        logger.info(
            "export_docx: %d placeholder(s) had no data and were left as "
            "'%s': %s",
            len(unresolved), pe_marker(), sorted(unresolved),
        )

    enable_update_fields_on_open(doc)

    pruned_acronyms = prune_unused_acronyms(doc)
    if pruned_acronyms:
        logger.info("export_docx: pruned %d unused acronym(s) from the Acronyms table", pruned_acronyms)

    output_path = export_dir / f"{project_name}_Phase1_ESA_DRAFT.docx"
    doc.save(str(output_path))
    logger.info("export_docx: wrote %s", output_path)
    return output_path
