"""
scripts/docx_helpers.py — Phase 1 ESA Report Generator

Low-level python-docx primitives used by scripts/export_docx.py to fill the
real Envicon_Phase_I_ESA_Report_TEMPLATE.docx template in place, rather than
re-generating a document from scratch (which is what loses all the template's
fixed structural content — cover page, live TOC, Revision History, Acronyms,
Signature Page — under a Pandoc-based approach).

These are the well-known python-docx gotchas, isolated here so they can be
unit-tested against small synthetic documents:

  - insert_paragraph_after   — python-docx only offers insert-BEFORE.
  - iter_all_paragraphs      — paragraphs live in the body, table cells, AND
                                headers/footers; placeholder substitution and
                                heading search must cover all of them.
  - build_heading_index      — locate template headings by style + normalized
                                text (numbering is typed into heading text,
                                not Word auto-numbering).
  - replace_placeholders     — Word fragments "{{Property Address}}" across
                                multiple runs (spellcheck/rsid splitting), so
                                a naive run.text.replace() silently misses
                                most placeholders.
  - find_table_by_header /
    append_row_like           — locate and grow one of the template's real
                                tables without losing its shading/borders.
  - enable_update_fields_on_open — the live TOC and "Page X of Y" footer
                                fields can't be computed by python-docx; ask
                                Word to refresh them the next time the file
                                is opened.
"""

from __future__ import annotations

import copy
import re
from dataclasses import dataclass

from docx.document import Document as DocumentObject
from docx.oxml.ns import qn
from docx.oxml.shared import OxmlElement
from docx.table import Table, _Cell, _Row
from docx.text.paragraph import Paragraph

_PLACEHOLDER_RE = re.compile(r"\{\{([^{}]+)\}\}")


# ---------------------------------------------------------------------------
# Traversal — every paragraph in the document, including tables & headers/footers
# ---------------------------------------------------------------------------

def iter_all_paragraphs(doc: DocumentObject):
    """
    Yield every Paragraph in the document: body paragraphs, every table cell's
    paragraphs (recursively, since a cell can itself contain a table), and
    every section's header/footer/first-page-header/first-page-footer.

    Both the cover page fields and the Signature Page table contain
    `{{placeholder}}` tokens, so placeholder substitution must traverse all
    of these or it will silently leave real placeholders unfilled.
    """
    yield from doc.paragraphs
    for table in doc.tables:
        yield from _iter_table_paragraphs(table)
    for section in doc.sections:
        for hf in (
            section.header,
            section.footer,
            section.first_page_header,
            section.first_page_footer,
        ):
            if hf is None:
                continue
            yield from hf.paragraphs
            for table in hf.tables:
                yield from _iter_table_paragraphs(table)


def _iter_table_paragraphs(table: Table):
    for row in table.rows:
        for cell in row.cells:
            yield from cell.paragraphs
            for nested in cell.tables:
                yield from _iter_table_paragraphs(nested)


# ---------------------------------------------------------------------------
# Insert-after — python-docx only provides insert_paragraph_before
# ---------------------------------------------------------------------------

def insert_paragraph_after(
    paragraph: Paragraph,
    text: str = "",
    style: str | None = None,
) -> Paragraph:
    """
    Insert a new paragraph immediately after *paragraph* and return it.

    python-docx exposes Paragraph.insert_paragraph_before() but nothing for
    "after" — the standard workaround is to build a bare <w:p> element and
    insert it via oxml addnext(), then wrap it back in a Paragraph object.

    Callers should NOT pass the heading's own style here — body content
    should use the template's "Normal" style (or "List Paragraph" for
    bullets), never the heading style, or the injected text will look like
    more headings. Falls back to no explicit style (Word default) if
    *style* isn't defined in the document, rather than raising, so a future
    template edit that renames/removes a style degrades gracefully instead
    of aborting the whole export.
    """
    new_p = OxmlElement("w:p")
    paragraph._p.addnext(new_p)
    new_para = Paragraph(new_p, paragraph._parent)
    if style is not None:
        try:
            new_para.style = style
        except KeyError:
            pass
    if text:
        new_para.add_run(text)
    return new_para


def remove_paragraph(paragraph: Paragraph) -> None:
    """Delete *paragraph* from the document entirely (removes its <w:p>
    element from its parent). Used to strip literal template scaffolding
    paragraphs (e.g. "» Writer note (delete before issue): ...") that should
    never reach an issued report."""
    p = paragraph._p
    p.getparent().remove(p)


def remove_paragraphs_matching(doc: DocumentObject, predicate) -> int:
    """Remove every body paragraph (not table cells/headers — the template's
    literal scaffolding paragraphs this is meant for only ever appear in the
    body) whose text satisfies *predicate*. Returns the count removed."""
    to_remove = [p for p in doc.paragraphs if predicate(p.text)]
    for p in to_remove:
        remove_paragraph(p)
    return len(to_remove)


# ---------------------------------------------------------------------------
# Body-order table lookup — for tables that share an identical header AND
# first-row signature (e.g. the template's Aerial and Sanborn tables both
# use "Year | Subject Property | Adjacent properties" with the same
# {{year}}/{{observation}} placeholder first row), find_table_by_header* has
# nothing left to disambiguate on. Locating by heading-proximity instead —
# these two tables are the only ones that immediately follow their own
# distinct heading in the template's document body order.
# ---------------------------------------------------------------------------

def remove_paragraphs_after_heading_matching(doc: DocumentObject, heading_para: Paragraph, predicate) -> int:
    """
    Remove every paragraph appearing after *heading_para* (in document body
    order) whose text satisfies *predicate*, stopping at the next
    Heading-styled paragraph — i.e. scoped to "this heading's own section"
    only. Unlike remove_paragraphs_matching (document-wide), this is for
    static boilerplate text that might coincidentally also appear as real
    content elsewhere in the document (e.g. a `{{Database Provider}}` token
    reused in more than one template section) — matching within a specific
    heading's section avoids deleting an unrelated paragraph that happens to
    share the same substring. Returns the count removed.
    """
    body = doc.element.body
    found_heading = False
    to_remove = []
    for child in body.iterchildren():
        if not found_heading:
            if child is heading_para._p:
                found_heading = True
            continue
        if child.tag == qn("w:p"):
            p_pr = child.find(qn("w:pPr"))
            style_val = None
            if p_pr is not None:
                p_style = p_pr.find(qn("w:pStyle"))
                if p_style is not None:
                    style_val = p_style.get(qn("w:val"))
            if style_val and style_val.startswith("Heading"):
                break
            text = "".join(t.text or "" for t in child.iter(qn("w:t")))
            if predicate(text):
                to_remove.append(child)
    for p in to_remove:
        p.getparent().remove(p)
    return len(to_remove)


def find_table_after_heading(doc: DocumentObject, heading_para: Paragraph) -> Table | None:
    """
    Return the first table appearing after *heading_para* in document body
    order, stopping (returning None) if another Heading-styled paragraph is
    reached first — i.e. only a table that sits directly under this specific
    heading (before the next one) counts as "this heading's table."
    """
    body = doc.element.body
    found_heading = False
    for child in body.iterchildren():
        if not found_heading:
            if child is heading_para._p:
                found_heading = True
            continue
        if child.tag == qn("w:tbl"):
            return Table(child, doc)
        if child.tag == qn("w:p"):
            p_pr = child.find(qn("w:pPr"))
            style_val = None
            if p_pr is not None:
                p_style = p_pr.find(qn("w:pStyle"))
                if p_style is not None:
                    style_val = p_style.get(qn("w:val"))
            if style_val and style_val.startswith("Heading"):
                return None
    return None


def insert_blocks_after(paragraph: Paragraph, blocks: list["Block"], normal_style: str = "Normal") -> Paragraph:
    """
    Insert a sequence of Block objects (see markdown_lite_to_blocks) after
    *paragraph*, in order, and return the last inserted paragraph (the new
    anchor for any further insertions).
    """
    anchor = paragraph
    for block in blocks:
        # The template defines "List Paragraph" (its bullets/numbering are
        # driven by numbering.xml via this style), not a "List Bullet" style.
        style = "List Paragraph" if block.is_bullet else normal_style
        anchor = insert_paragraph_after(anchor, style=style)
        _add_runs_with_bold(anchor, block.text)
    return anchor


def _add_runs_with_bold(paragraph: Paragraph, text: str) -> None:
    """Add runs to *paragraph*, splitting on **bold** markers into bold runs."""
    parts = re.split(r"(\*\*[^*]+\*\*)", text)
    for part in parts:
        if not part:
            continue
        if part.startswith("**") and part.endswith("**") and len(part) > 4:
            run = paragraph.add_run(part[2:-2])
            run.bold = True
        else:
            paragraph.add_run(part)


# ---------------------------------------------------------------------------
# Minimal markdown -> block list (Writer prose is markdown-lite: paragraphs,
# "- " / "* " bullets, **bold**). Tables are deliberately NOT handled here —
# the 6 real tables are built programmatically from pipeline data, not from
# Writer-emitted markdown tables.
# ---------------------------------------------------------------------------

@dataclass
class Block:
    text: str
    is_bullet: bool = False


def markdown_lite_to_blocks(text: str) -> list[Block]:
    blocks: list[Block] = []
    for raw_line in text.splitlines():
        line = raw_line.strip()
        if not line:
            continue
        if line.startswith(("- ", "* ")):
            blocks.append(Block(text=line[2:].strip(), is_bullet=True))
        else:
            blocks.append(Block(text=line))
    return blocks


# ---------------------------------------------------------------------------
# Heading lookup
# ---------------------------------------------------------------------------

def normalize_heading(text: str) -> str:
    """Collapse whitespace and casefold, for tolerant heading-text matching."""
    return re.sub(r"\s+", " ", text).strip().casefold()


def build_heading_index(doc: DocumentObject) -> dict[str, Paragraph]:
    """
    Map normalized heading text -> its Paragraph, for every paragraph whose
    style name starts with "Heading" (confirmed present in the real template:
    w:pStyle val="Heading1/2/3", not direct formatting).

    Section numbering in the template is typed directly into the heading
    text ("1.0 Introduction", "5.2.1 Aerial Photographs") rather than driven
    by Word auto-numbering, so matching on normalized full text is reliable
    and doesn't require reproducing a numbering scheme.
    """
    index: dict[str, Paragraph] = {}
    for para in doc.paragraphs:
        style_name = para.style.name if para.style is not None else ""
        if not style_name.startswith("Heading"):
            continue
        key = normalize_heading(para.text)
        if key:
            index[key] = para
    return index


# ---------------------------------------------------------------------------
# Placeholder substitution — multi-run-aware
# ---------------------------------------------------------------------------

def replace_placeholders(
    doc: DocumentObject,
    mapping: dict[str, str],
    unresolved_marker: str | None = None,
) -> set[str]:
    """
    Replace every `{{Name}}` token anywhere in the document with
    mapping[normalized_name] if present.

    Word frequently splits a single `{{Property Address}}` token across
    multiple runs (spellcheck / revision-tracking boundaries), so this
    cannot be a simple per-run string replace. This implementation avoids
    the classic bug of re-reading/re-slicing run.text after it has already
    been mutated (which corrupts results when two placeholders share a run,
    or a placeholder spans multiple runs) by working entirely from an
    immutable per-paragraph snapshot:

      1. Snapshot every run's original text and its [start, end) offset in
         the paragraph's concatenated text — never read from these runs
         again during this paragraph's processing.
      2. Walk the matches left to right, accumulating output pieces per run
         index in `new_run_pieces`: unchanged spans keep their original
         per-run attribution (preserving that text's formatting), and each
         match's replacement text is appended entirely to the first run the
         match touches.
      3. Only after all matches are processed does each run's text get
         assigned exactly once (`run.text = "".join(pieces)`), and only runs
         that received a replacement have their placeholder formatting
         (highlight + color) stripped.

    Placeholders with no entry in *mapping* are left as *unresolved_marker*
    (if given) instead of the raw `{{...}}` token, and their normalized
    names are returned so callers can log/assert on them. Never silently
    blank an unresolved placeholder.
    """
    unresolved: set[str] = set()

    for paragraph in iter_all_paragraphs(doc):
        runs = paragraph.runs
        if not runs:
            continue
        orig_texts = [r.text for r in runs]
        full_text = "".join(orig_texts)
        if "{{" not in full_text:
            continue

        matches = list(_PLACEHOLDER_RE.finditer(full_text))
        if not matches:
            continue

        run_bounds: list[tuple[int, int]] = []
        offset = 0
        for t in orig_texts:
            run_bounds.append((offset, offset + len(t)))
            offset += len(t)

        new_run_pieces: list[list[str]] = [[] for _ in runs]
        dehighlight: set[int] = set()

        def _emit_original_span(start: int, end: int) -> None:
            if start >= end:
                return
            for i, (rs, re_) in enumerate(run_bounds):
                seg_start, seg_end = max(start, rs), min(end, re_)
                if seg_start < seg_end:
                    new_run_pieces[i].append(orig_texts[i][seg_start - rs:seg_end - rs])

        cursor = 0
        for match in matches:
            m_start, m_end = match.span()
            name = match.group(1).strip()
            key = name.casefold()
            replacement = mapping.get(key)
            if replacement is None:
                unresolved.add(name)
                replacement = unresolved_marker if unresolved_marker is not None else match.group(0)

            _emit_original_span(cursor, m_start)

            first_idx = next(
                (i for i, (rs, re_) in enumerate(run_bounds) if rs < m_end and re_ > m_start),
                None,
            )
            if first_idx is not None:
                new_run_pieces[first_idx].append(replacement)
                dehighlight.add(first_idx)
            cursor = m_end

        _emit_original_span(cursor, len(full_text))

        for i, run in enumerate(runs):
            run.text = "".join(new_run_pieces[i])
            if i in dehighlight:
                _strip_placeholder_formatting(run)

    return unresolved


def _strip_placeholder_formatting(run) -> None:
    """Clear the template's brown-text/yellow-highlight placeholder styling
    from *run* so substituted real data reads as normal body text."""
    run.font.highlight_color = None
    try:
        run.font.color.rgb = None
    except Exception:  # noqa: BLE001 - color may not be directly settable in all cases
        pass


# ---------------------------------------------------------------------------
# Table lookup + row cloning
# ---------------------------------------------------------------------------

def find_table_by_header(doc: DocumentObject, signature: tuple[str, ...]) -> Table | None:
    """
    Return the first table in *doc* whose first row's cell texts match
    *signature* (normalized, order-preserving). Searching by header-cell
    text is robust to the table's position in the document, unlike indexing
    by table position.
    """
    normalized_sig = tuple(normalize_heading(s) for s in signature)
    for table in doc.tables:
        if not table.rows:
            continue
        header_cells = tuple(normalize_heading(c.text) for c in table.rows[0].cells)
        if header_cells == normalized_sig:
            return table
    return None


def find_table_by_header_and_first_row(
    doc: DocumentObject,
    signature: tuple[str, ...],
    first_data_cell_text: str,
) -> Table | None:
    """
    Like find_table_by_header, but disambiguates between multiple tables
    that share an identical header row (e.g. the template's Federal and
    State/Tribal/Local EDR radius tables both use
    "List | Search radius | On Subject Property | Listings in radius | REC
    relative to Subject Property") by additionally checking that the first
    data row's first cell matches *first_data_cell_text* (normalized,
    case-insensitive) — e.g. "NPL" for the Federal table, "SHWS" for the
    State table.
    """
    normalized_sig = tuple(normalize_heading(s) for s in signature)
    target_first_cell = normalize_heading(first_data_cell_text)
    for table in doc.tables:
        if len(table.rows) < 2:
            continue
        header_cells = tuple(normalize_heading(c.text) for c in table.rows[0].cells)
        if header_cells != normalized_sig:
            continue
        if normalize_heading(table.rows[1].cells[0].text) == target_first_cell:
            return table
    return None


def set_cell_text(cell: _Cell, text: str) -> None:
    """
    Overwrite *cell*'s content with a single run of *text*, preserving the
    cell's own paragraph/run formatting (font, shading) rather than
    inserting a brand-new default-styled paragraph.
    """
    if cell.paragraphs:
        para = cell.paragraphs[0]
        for run in list(para.runs):
            run.text = ""
        if para.runs:
            para.runs[0].text = text
        else:
            para.add_run(text)
        # Remove any extra paragraphs beyond the first so old content
        # doesn't linger below the new text.
        for extra in cell.paragraphs[1:]:
            extra._p.getparent().remove(extra._p)
    else:
        cell.add_paragraph(text)


def append_row_like(table: Table, template_row_index: int, values: list[str]) -> _Row:
    """
    Clone the table row at *template_row_index* (deep-copies its <w:tr>, so
    the new row inherits whatever direct cell shading/borders/fonts the
    template row has — the template has no named table style, so this is
    the only way to preserve its look) and overwrite its cell text with
    *values* (by position; extra/missing values are ignored/left blank).
    """
    src_tr = table.rows[template_row_index]._tr
    new_tr = copy.deepcopy(src_tr)
    table._tbl.append(new_tr)
    new_row = table.rows[-1]
    for cell, value in zip(new_row.cells, values):
        set_cell_text(cell, value)
    return new_row


# ---------------------------------------------------------------------------
# Live fields (TOC, Page X of Y) — can't be computed here; ask Word to refresh
# ---------------------------------------------------------------------------

def enable_update_fields_on_open(doc: DocumentObject) -> None:
    """
    Set <w:updateFields w:val="true"/> in document settings so Word
    recalculates the live Table-of-Contents and "Page X of Y" footer fields
    the next time this file is opened. python-docx has no pagination engine,
    so these fields cannot be computed at export time — this is the standard
    way to defer that computation to Word itself without requiring Word (or
    LibreOffice) to be installed on the machine running the pipeline.
    """
    settings = doc.settings.element
    el = settings.find(qn("w:updateFields"))
    if el is None:
        el = OxmlElement("w:updateFields")
        settings.append(el)
    el.set(qn("w:val"), "true")
