"""
Unit tests for the WS7 export fixes in scripts/export_docx.py (+ the
docx_helpers.py primitives they depend on): stripping the template's literal
"Writer note" paragraphs, removing its duplicate placeholder References
list, and populating the aerial/Sanborn/city-directory historical tables.

Builds small synthetic python-docx documents rather than depending on the
real Envicon template binary, so these run standalone and fast.
"""

from __future__ import annotations

import json
from pathlib import Path

from docx import Document

from scripts.docx_helpers import build_heading_index, find_table_after_heading, find_table_by_header
from scripts.export_docx import (
    populate_historical_tables,
    prune_unused_acronyms,
    strip_duplicate_references_boilerplate,
    strip_writer_notes,
)


def _make_doc_with_writer_notes():
    doc = Document()
    doc.add_heading("Revision History", level=1)
    doc.add_paragraph("» Writer note (delete before issue): Some clients want this page removed.")
    doc.add_paragraph("Normal body content that should survive.")
    doc.add_heading("List of Acronyms", level=1)
    doc.add_paragraph("» Writer note (delete before issue): Keep only acronyms actually used.")
    return doc


class TestStripWriterNotes:
    def test_removes_all_writer_note_paragraphs(self):
        doc = _make_doc_with_writer_notes()
        removed = strip_writer_notes(doc)
        assert removed == 2
        remaining_text = "\n".join(p.text for p in doc.paragraphs)
        assert "Writer note" not in remaining_text
        assert "Normal body content that should survive." in remaining_text

    def test_no_op_when_none_present(self):
        doc = Document()
        doc.add_paragraph("Nothing to strip here.")
        assert strip_writer_notes(doc) == 0


def _make_doc_with_references_section(inject_real_references: bool = True):
    doc = Document()
    doc.add_heading("12.0 References", level=1)
    if inject_real_references:
        doc.add_paragraph("ASTM E1527-21 and the EDR Radius Map Report, dated Jan 30, 2024.")
    doc.add_paragraph("The following sources were reviewed as part of this assessment:")
    doc.add_paragraph(
        "ASTM International. Standard Practice for Environmental Site Assessments: Phase I "
        "Environmental Site Assessment Process. Designation E1527-21."
    )
    doc.add_paragraph("{{Database Provider}} Radius Map Report, {{Property Address}}, {{Database Report Date}}.")
    doc.add_paragraph("{{Database Provider}} Certified Sanborn Map Report, {{Property Address}}, {{date}}.")
    doc.add_paragraph("{{Database Provider}} City Directory Package, {{Property Address}}, {{date}}.")
    doc.add_paragraph("{{Database Provider}} Aerial Photographs, {{Property Address}}, {{date}}.")
    doc.add_paragraph("{{Prior report title, author, and date}}")
    doc.add_paragraph("{{Agency records reviewed (NYSDEC, county, municipal), with dates}}")
    doc.add_heading("Signature Page", level=1)
    return doc


class TestStripDuplicateReferencesBoilerplate:
    def test_removes_boilerplate_when_real_references_were_drafted(self):
        doc = _make_doc_with_references_section(inject_real_references=True)
        headings = build_heading_index(doc)
        removed = strip_duplicate_references_boilerplate(doc, {"12.0 references": "real content"}, headings)
        assert removed == 8
        remaining_text = "\n".join(p.text for p in doc.paragraphs)
        assert "{{Database Provider}}" not in remaining_text
        assert "The following sources were reviewed" not in remaining_text
        # The real, drafted References content survives untouched.
        assert "EDR Radius Map Report, dated Jan 30, 2024." in remaining_text

    def test_leaves_boilerplate_untouched_when_nothing_was_drafted(self):
        doc = _make_doc_with_references_section(inject_real_references=False)
        headings = build_heading_index(doc)
        removed = strip_duplicate_references_boilerplate(doc, {}, headings)  # no "12.0 references" key
        assert removed == 0
        remaining_text = "\n".join(p.text for p in doc.paragraphs)
        assert "{{Database Provider}}" in remaining_text

    def test_does_not_delete_a_shared_token_used_elsewhere_in_the_document(self):
        # Regression test: the real Envicon template's Section 5.3 intro
        # paragraph ALSO contains "{{Database Provider}} Radius Map Report"
        # (legitimate content, not the duplicate References list) — a
        # document-wide substring match would incorrectly delete it too.
        doc = _make_doc_with_references_section(inject_real_references=True)
        doc.add_heading("5.3 Standard Federal, State, Tribal, and Local Records", level=2)
        doc.add_paragraph(
            "Envicon reviewed the {{Database Provider}} Radius Map Report dated "
            "{{Database Report Date}} using the search distances specified in ASTM E1527-21."
        )
        headings = build_heading_index(doc)

        strip_duplicate_references_boilerplate(doc, {"12.0 references": "real content"}, headings)

        remaining_text = "\n".join(p.text for p in doc.paragraphs)
        assert "Envicon reviewed the {{Database Provider}} Radius Map Report" in remaining_text


class TestFindTableAfterHeading:
    def test_distinguishes_identical_tables_by_heading_proximity(self):
        doc = Document()
        headings = {}
        doc.add_heading("5.2.1 Aerial Photographs", level=3)
        t1 = doc.add_table(rows=2, cols=3)
        t1.rows[0].cells[0].text = "Year"
        t1.rows[1].cells[0].text = "{{year}}"
        doc.add_heading("5.2.2 Sanborn Fire Insurance Maps", level=3)
        t2 = doc.add_table(rows=2, cols=3)
        t2.rows[0].cells[0].text = "Year"
        t2.rows[1].cells[0].text = "{{year}}"

        headings = build_heading_index(doc)
        aerial_heading = headings["5.2.1 aerial photographs"]
        sanborn_heading = headings["5.2.2 sanborn fire insurance maps"]

        found_aerial = find_table_after_heading(doc, aerial_heading)
        found_sanborn = find_table_after_heading(doc, sanborn_heading)

        assert found_aerial is not None and found_sanborn is not None
        assert found_aerial._tbl is not found_sanborn._tbl

    def test_returns_none_if_next_heading_reached_before_a_table(self):
        doc = Document()
        doc.add_heading("5.2.1 Aerial Photographs", level=3)
        doc.add_paragraph("No table under this heading.")
        doc.add_heading("5.2.2 Sanborn Fire Insurance Maps", level=3)
        doc.add_table(rows=2, cols=3)

        headings = build_heading_index(doc)
        assert find_table_after_heading(doc, headings["5.2.1 aerial photographs"]) is None


def _make_doc_for_historical_tables():
    doc = Document()
    doc.add_heading("5.2.1 Aerial Photographs", level=3)
    aerial = doc.add_table(rows=2, cols=3)
    for i, h in enumerate(("Year", "Subject Property", "Adjacent properties")):
        aerial.rows[0].cells[i].text = h
    for i, v in enumerate(("{{year}}", "{{observation}}", "{{observation}}")):
        aerial.rows[1].cells[i].text = v

    doc.add_heading("5.2.2 Sanborn Fire Insurance Maps", level=3)
    sanborn = doc.add_table(rows=2, cols=3)
    for i, h in enumerate(("Year", "Subject Property", "Adjacent properties")):
        sanborn.rows[0].cells[i].text = h
    for i, v in enumerate(("{{year}}", "{{observation}}", "{{observation}}")):
        sanborn.rows[1].cells[i].text = v

    doc.add_heading("5.2.3 City and Street Directories", level=3)
    directory = doc.add_table(rows=2, cols=3)
    for i, h in enumerate(("Year", "Address", "Occupant")):
        directory.rows[0].cells[i].text = h
    for i, v in enumerate(("{{year}}", "{{address}}", "{{occupant}}")):
        directory.rows[1].cells[i].text = v

    return doc


class TestPopulateHistoricalTables:
    def test_fills_aerial_and_leaves_sanborn_untouched_when_no_data(self, tmp_path: Path):
        project = tmp_path / "TestProject"
        (project / "Historical_Records").mkdir(parents=True)
        (project / "Historical_Records" / "historical_tables.json").write_text(
            json.dumps({
                "aerial": [
                    {"year": "1950", "subject_property": "Vacant lot", "adjacent_properties": "Industrial"},
                    {"year": "1986", "subject_property": "Factory building", "adjacent_properties": "Rail spur"},
                ],
                "sanborn": [],
                "city_directory": [],
            }),
            encoding="utf-8",
        )
        doc = _make_doc_for_historical_tables()
        headings = build_heading_index(doc)

        populate_historical_tables(doc, project, headings)

        aerial_table = find_table_after_heading(doc, headings["5.2.1 aerial photographs"])
        assert len(aerial_table.rows) == 3  # header + 2 real rows
        assert aerial_table.rows[1].cells[0].text == "1950"
        assert aerial_table.rows[1].cells[1].text == "Vacant lot"
        assert aerial_table.rows[2].cells[0].text == "1986"

        sanborn_table = find_table_after_heading(doc, headings["5.2.2 sanborn fire insurance maps"])
        assert sanborn_table.rows[1].cells[0].text == "{{year}}"  # untouched — no data

    def test_no_op_when_json_file_absent(self, tmp_path: Path):
        project = tmp_path / "TestProject"
        project.mkdir()
        doc = _make_doc_for_historical_tables()
        headings = build_heading_index(doc)
        populate_historical_tables(doc, project, headings)  # must not raise
        aerial_table = find_table_after_heading(doc, headings["5.2.1 aerial photographs"])
        assert aerial_table.rows[1].cells[0].text == "{{year}}"

    def test_city_directory_uses_address_occupant_columns(self, tmp_path: Path):
        project = tmp_path / "TestProject"
        (project / "Historical_Records").mkdir(parents=True)
        (project / "Historical_Records" / "historical_tables.json").write_text(
            json.dumps({
                "aerial": [], "sanborn": [],
                "city_directory": [{"year": "1965", "address": "631 Northland Ave", "occupant": "Example Co."}],
            }),
            encoding="utf-8",
        )
        doc = _make_doc_for_historical_tables()
        headings = build_heading_index(doc)
        populate_historical_tables(doc, project, headings)
        directory_table = find_table_after_heading(doc, headings["5.2.3 city and street directories"])
        assert directory_table.rows[1].cells[1].text == "631 Northland Ave"
        assert directory_table.rows[1].cells[2].text == "Example Co."


class TestPruneUnusedAcronyms:
    def test_removes_rows_for_acronyms_never_mentioned_elsewhere(self):
        doc = Document()
        doc.add_heading("List of Acronyms", level=1)
        table = doc.add_table(rows=3, cols=2)
        table.rows[0].cells[0].text, table.rows[0].cells[1].text = "Acronym", "Definition"
        table.rows[1].cells[0].text, table.rows[1].cells[1].text = "AAI", "All Appropriate Inquiries"
        table.rows[2].cells[0].text, table.rows[2].cells[1].text = "PFAS", "Per- and polyfluoroalkyl substances"
        doc.add_paragraph("This assessment was conducted per AAI standards.")
        # PFAS is never mentioned in the body.

        removed = prune_unused_acronyms(doc)

        assert removed == 1
        remaining = [row.cells[0].text for row in table.rows[1:]]
        assert remaining == ["AAI"]

    def test_no_op_when_no_acronym_table_present(self):
        doc = Document()
        doc.add_paragraph("No acronyms table here.")
        assert prune_unused_acronyms(doc) == 0
